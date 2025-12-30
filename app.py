import streamlit as st
import os
import logging
import warnings
from PIL import Image
import numpy as np
import io
from docx import Document

# ================= 1. 环境配置 =================
os.environ["DISABLE_MODEL_SOURCE_CHECK"] = "1"
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
logging.getLogger("ppocr").setLevel(logging.ERROR)
warnings.filterwarnings("ignore")

from paddleocr import PaddleOCR

# ================= 2. 核心逻辑 =================

@st.cache_resource
def load_ocr_model():
    # 初始化
    return PaddleOCR(use_angle_cls=True, lang="ch", enable_mkldnn=False)

def parse_ocr_result(data):
    """
    【核心修复】针对新版 PaddleOCR 返回的字典结构进行解析
    """
    texts = []
    
    # 1. 如果数据直接是字典 (包含 rec_texts 键)
    if isinstance(data, dict) and 'rec_texts' in data:
        return data['rec_texts']
        
    # 2. 如果数据是列表 (这是日志显示的结构: [{'rec_texts': [...]}])
    if isinstance(data, list):
        for item in data:
            # 情况 A: 列表里的元素是字典，且包含 rec_texts
            if isinstance(item, dict) and 'rec_texts' in item:
                texts.extend(item['rec_texts'])
                
            # 情况 B: 旧版标准结构 [box, (text, score)]
            elif isinstance(item, tuple) and len(item) == 2 and isinstance(item[0], str):
                texts.append(item[0])
                
            # 情况 C: 旧版列表结构 [text, score]
            elif isinstance(item, list) and len(item) == 2 and isinstance(item[0], str) and isinstance(item[1], (float, int)):
                texts.append(item[0])
                
            # 情况 D: 嵌套列表，继续递归
            elif isinstance(item, list):
                texts.extend(parse_ocr_result(item))
                
    return texts

def process_ocr(image, ocr_model):
    # 确保图片是 uint8 格式
    img_array = np.array(image.convert('RGB'), dtype='uint8')
    
    print(">>> 开始识别...")
    try:
        # 直接调用
        result = ocr_model.ocr(img_array)
    except Exception as e:
        return f"识别出错: {e}"

    if result is None:
        return ""

    # 使用新的解析函数
    extracted_texts = parse_ocr_result(result)
    
    # 打印调试信息
    print(f">>> 成功提取行数: {len(extracted_texts)}")
    if len(extracted_texts) > 0:
        print(f">>> 第一行内容: {extracted_texts[0]}")
    
    return "\n".join(extracted_texts)

def create_word_doc(text_content):
    doc = Document()
    # doc.add_heading('OCR 识别结果', 0)
    for line in text_content.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ================= 3. 页面 UI =================

st.set_page_config(page_title="OCR转Word", page_icon="📝")
st.title("📝 图片转 Word (适配新版)")

with st.spinner('加载 AI 模型...'):
    ocr_engine = load_ocr_model()

uploaded_file = st.file_uploader("上传图片", type=['png', 'jpg', 'jpeg', 'bmp'])

if uploaded_file:
    image = Image.open(uploaded_file)
    st.image(image, caption='预览', width="stretch")
    
    if st.button("开始转换", type="primary", width="stretch"):
        with st.spinner('正在识别文字...'):
            text = process_ocr(image, ocr_engine)
            
            if text and text.strip():
                st.success(f"✅ 识别成功！共 {len(text)} 字。")
                st.text_area("识别结果", text, height=300)
                st.download_button(
                    "📥 下载 Word 文档", 
                    create_word_doc(text), 
                    "ocr_result.docx", 
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    width="stretch"
                )
            else:
                st.error("❌ 依然没有提取到文字。")